import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_inline_formatting(paragraph, text):
    """
    Parses simple markdown inline formatting (bold: **text**) and adds runs to paragraph.
    """
    # Regex split to isolate bold sections
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            paragraph.add_run(part)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set standard page margins (1 inch / 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set base directory to resolve relative images
    base_dir = os.path.dirname(os.path.abspath(md_path))

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            continue

        # Headings
        if stripped.startswith('# '):
            heading_text = stripped[2:]
            h = doc.add_heading(level=1)
            parse_inline_formatting(h, heading_text)
            continue
        elif stripped.startswith('## '):
            heading_text = stripped[3:]
            h = doc.add_heading(level=2)
            parse_inline_formatting(h, heading_text)
            continue
        elif stripped.startswith('### '):
            heading_text = stripped[4:]
            h = doc.add_heading(level=3)
            parse_inline_formatting(h, heading_text)
            continue
        elif stripped.startswith('#### '):
            heading_text = stripped[5:]
            h = doc.add_heading(level=6) # Sub-item
            parse_inline_formatting(h, heading_text)
            continue
        
        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('__________________________________________________')
            run.font.color.rgb = RGBColor(180, 180, 180)
            continue

        # LaTeX block formulas $$ ... $$
        if stripped.startswith('$$') and stripped.endswith('$$'):
            formula = stripped[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            continue

        # Image check: ![alt](path)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2)
            # Resolve absolute path
            img_abs_path = os.path.abspath(os.path.join(base_dir, img_rel_path.replace('/', '\\')))
            if os.path.exists(img_abs_path):
                # Add picture
                doc.add_picture(img_abs_path, width=Inches(5.8))
                # Add caption below picture
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Hình: {alt_text}")
                run.italic = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(100, 100, 100)
            else:
                p = doc.add_paragraph()
                p.add_run(f"[Hình ảnh không tìm thấy: {img_rel_path}]").bold = True
            continue

        # Bullet list item
        if stripped.startswith('* ') or stripped.startswith('- '):
            list_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, list_text)
            continue

        # Numbered list item
        numbered_match = re.match(r'^\d+\.\s+(.*)$', stripped)
        if numbered_match:
            list_text = numbered_match.group(1)
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, list_text)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped)

    # Save file
    doc.save(docx_path)
    print(f"Successfully converted {md_path} -> {docx_path}")

# Run conversion
if __name__ == "__main__":
    docs_dir = r"d:\My_projects\XAI_KLTN\docs"
    files_to_convert = [
        "vi/1_Tong_Quan_De_Tai.md",
        "vi/2_Tai_Lieu_He_Thong_Web.md",
        "vi/3_Tai_Lieu_Mo_Hinh_AI.md",
        "vi/4_Bao_Cao_Thuc_Nghiem_So_Bo.md",
        "vi/5_Phat_Hien_Nghien_Cuu_Moi_Phase_2.md",
        "vi/Research_Methodology_and_Statistical_Design.md",
        "en/1_Project_Overview.md",
        "en/2_Web_System_Documentation.md",
        "en/3_AI_Model_Documentation.md",
        "en/4_Pilot_Experiment_Report.md",
        "en/5_New_Research_Findings_Phase_2.md",
        "en/Research_Methodology_and_Statistical_Design.md",
        "test/local_test_analysis.md"
    ]

    for filename in files_to_convert:
        md_file = os.path.join(docs_dir, filename)
        docx_file = md_file.replace(".md", ".docx")
        convert_md_to_docx(md_file, docx_file)
